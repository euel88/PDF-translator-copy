"""
Word 문서 핸들러 (.docx, .doc)
python-docx 라이브러리를 사용하여 Word 문서 번역

개선된 기능:
- Run 단위 스타일 완벽 보존 (볼드, 이탤릭, 색상, 폰트 등)
- 페이지 레이아웃 최대한 보존
- 대용량 문서 청크 처리
- 번역 텍스트 길이 변화에 대응하는 동적 조정
"""
from typing import Optional, List, Dict, Any, Tuple
from pathlib import Path
from dataclasses import dataclass
import copy

from pdf2zh.handlers.base import BaseDocumentHandler, DocumentResult


@dataclass
class RunInfo:
    """Run 정보 저장"""
    text: str
    # 스타일 정보
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None
    strike: Optional[bool] = None
    font_name: Optional[str] = None
    font_size: Optional[int] = None
    font_color_rgb: Optional[Tuple[int, int, int]] = None
    highlight_color: Optional[str] = None


@dataclass
class TextUnit:
    """번역 단위 (Run 또는 단락)"""
    location_type: str  # 'paragraph', 'table', 'header', 'footer'
    location_ids: tuple  # 위치 식별자
    runs_info: List[RunInfo]  # 각 run의 정보
    full_text: str  # 전체 텍스트 (번역용)


class WordHandler(BaseDocumentHandler):
    """Word 문서 핸들러 - 개선된 스타일 보존"""

    SUPPORTED_EXTENSIONS = ['.docx', '.doc']

    # 대용량 처리 설정
    BATCH_SIZE_SMALL = 30   # 짧은 텍스트
    BATCH_SIZE_MEDIUM = 20  # 중간 텍스트
    BATCH_SIZE_LARGE = 10   # 긴 텍스트

    def convert(
        self,
        input_path: str,
        output_path: Optional[str] = None
    ) -> DocumentResult:
        """
        Word 문서 번역 - Run 단위 스타일 보존

        Args:
            input_path: 입력 파일 경로
            output_path: 출력 파일 경로

        Returns:
            DocumentResult: 변환 결과
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_UNDERLINE
        except ImportError:
            return DocumentResult(
                success=False,
                error="python-docx 라이브러리가 설치되지 않았습니다. 'pip install python-docx'를 실행하세요."
            )

        if output_path is None:
            output_path = self.generate_output_path(input_path)

        try:
            self.log(f"Word 문서 열기: {input_path}")
            doc = Document(input_path)

            # 1. 텍스트 단위 수집 (Run 정보 포함)
            text_units = self._collect_text_units(doc)
            total_count = len(text_units)
            self.log(f"번역할 텍스트 단위: {total_count}개")

            if total_count == 0:
                doc.save(output_path)
                return DocumentResult(
                    output_path=output_path,
                    success=True,
                    translated_count=0,
                    total_count=0
                )

            # 2. 동적 배치 크기로 번역
            self.log("번역 중...")
            translations = self._translate_with_dynamic_batch(text_units)

            # 3. 번역 결과 적용 (스타일 보존)
            translated_count = self._apply_translations(doc, text_units, translations)

            # 4. 저장
            self.log(f"저장 중: {output_path}")
            doc.save(output_path)

            return DocumentResult(
                output_path=output_path,
                success=True,
                translated_count=translated_count,
                total_count=total_count
            )

        except Exception as e:
            import traceback
            return DocumentResult(
                success=False,
                error=f"{str(e)}\n{traceback.format_exc()}"
            )

    def _collect_text_units(self, doc) -> List[TextUnit]:
        """텍스트 단위 수집 - Run 정보 포함 (모든 텍스트 요소 포함)"""
        units = []
        from docx.oxml.ns import qn

        # 수집 통계
        stats = {
            'paragraph': 0, 'table': 0, 'header': 0, 'footer': 0,
            'footnote': 0, 'endnote': 0, 'textbox': 0, 'sdt': 0
        }

        # 1. 본문 단락 수집
        for para_idx, para in enumerate(doc.paragraphs):
            unit = self._extract_paragraph_unit(para, 'paragraph', (para_idx,))
            if unit:
                units.append(unit)
                stats['paragraph'] += 1

        # 2. 표 수집 (중첩 표 포함)
        for table_idx, table in enumerate(doc.tables):
            self._collect_table_text(table, table_idx, units, stats)

        # 3. 헤더/푸터 수집 (모든 유형)
        for section_idx, section in enumerate(doc.sections):
            # 모든 헤더 유형 수집
            for header_type in ['header', 'first_page_header', 'even_page_header']:
                try:
                    header = getattr(section, header_type, None)
                    if header and header.is_linked_to_previous is False:
                        for para_idx, para in enumerate(header.paragraphs):
                            unit = self._extract_paragraph_unit(
                                para, 'header', (section_idx, header_type, para_idx)
                            )
                            if unit:
                                units.append(unit)
                                stats['header'] += 1
                        # 헤더 내 표
                        for tbl_idx, table in enumerate(header.tables):
                            self._collect_table_text(
                                table, f"header_{section_idx}_{tbl_idx}",
                                units, stats, location_prefix='header_table'
                            )
                except Exception:
                    pass

            # 모든 푸터 유형 수집
            for footer_type in ['footer', 'first_page_footer', 'even_page_footer']:
                try:
                    footer = getattr(section, footer_type, None)
                    if footer and footer.is_linked_to_previous is False:
                        for para_idx, para in enumerate(footer.paragraphs):
                            unit = self._extract_paragraph_unit(
                                para, 'footer', (section_idx, footer_type, para_idx)
                            )
                            if unit:
                                units.append(unit)
                                stats['footer'] += 1
                        # 푸터 내 표
                        for tbl_idx, table in enumerate(footer.tables):
                            self._collect_table_text(
                                table, f"footer_{section_idx}_{tbl_idx}",
                                units, stats, location_prefix='footer_table'
                            )
                except Exception:
                    pass

        # 4. 콘텐츠 컨트롤 (SDT) 수집 - 많은 Word 문서에서 사용
        try:
            body = doc.element.body
            for sdt_idx, sdt in enumerate(body.iter(qn('w:sdt'))):
                # SDT 내의 모든 텍스트 수집
                sdt_content = sdt.find(qn('w:sdtContent'))
                if sdt_content is not None:
                    for para_idx, para_elem in enumerate(sdt_content.findall(qn('w:p'))):
                        text = ''.join(t.text or '' for t in para_elem.iter(qn('w:t')))
                        if text.strip():
                            units.append(TextUnit(
                                location_type='sdt',
                                location_ids=(sdt_idx, para_idx),
                                runs_info=[RunInfo(text=text)],
                                full_text=text
                            ))
                            stats['sdt'] += 1
        except Exception:
            pass

        # 5. 각주(Footnotes) 수집
        try:
            if hasattr(doc, 'part') and doc.part is not None:
                footnotes_part = None
                for rel in doc.part.rels.values():
                    if 'footnotes' in rel.reltype:
                        footnotes_part = rel.target_part
                        break

                if footnotes_part is not None:
                    footnotes_xml = footnotes_part.element
                    for fn_idx, footnote in enumerate(footnotes_xml.findall(qn('w:footnote'))):
                        for para_idx, para_elem in enumerate(footnote.findall(qn('w:p'))):
                            text = ''.join(t.text or '' for t in para_elem.iter(qn('w:t')))
                            if text.strip():
                                units.append(TextUnit(
                                    location_type='footnote',
                                    location_ids=(fn_idx, para_idx),
                                    runs_info=[RunInfo(text=text)],
                                    full_text=text
                                ))
                                stats['footnote'] += 1
        except Exception:
            pass

        # 6. 미주(Endnotes) 수집
        try:
            if hasattr(doc, 'part') and doc.part is not None:
                endnotes_part = None
                for rel in doc.part.rels.values():
                    if 'endnotes' in rel.reltype:
                        endnotes_part = rel.target_part
                        break

                if endnotes_part is not None:
                    endnotes_xml = endnotes_part.element
                    for en_idx, endnote in enumerate(endnotes_xml.findall(qn('w:endnote'))):
                        for para_idx, para_elem in enumerate(endnote.findall(qn('w:p'))):
                            text = ''.join(t.text or '' for t in para_elem.iter(qn('w:t')))
                            if text.strip():
                                units.append(TextUnit(
                                    location_type='endnote',
                                    location_ids=(en_idx, para_idx),
                                    runs_info=[RunInfo(text=text)],
                                    full_text=text
                                ))
                                stats['endnote'] += 1
        except Exception:
            pass

        # 7. 텍스트 박스 (도형 내 텍스트) 수집
        try:
            body = doc.element.body
            for txbx_idx, txbx in enumerate(body.iter(qn('w:txbxContent'))):
                for para_idx, para_elem in enumerate(txbx.findall(qn('w:p'))):
                    text = ''.join(t.text or '' for t in para_elem.iter(qn('w:t')))
                    if text.strip():
                        units.append(TextUnit(
                            location_type='textbox',
                            location_ids=(txbx_idx, para_idx),
                            runs_info=[RunInfo(text=text)],
                            full_text=text
                        ))
                        stats['textbox'] += 1
        except Exception:
            pass

        # 8. 주석(Comments) 수집
        stats['comment'] = 0
        try:
            if hasattr(doc, 'part') and doc.part is not None:
                comments_part = None
                for rel in doc.part.rels.values():
                    if 'comments' in rel.reltype:
                        comments_part = rel.target_part
                        break

                if comments_part is not None:
                    comments_xml = comments_part.element
                    for cmt_idx, comment in enumerate(comments_xml.findall(qn('w:comment'))):
                        for para_idx, para_elem in enumerate(comment.findall(qn('w:p'))):
                            text = ''.join(t.text or '' for t in para_elem.iter(qn('w:t')))
                            if text.strip():
                                units.append(TextUnit(
                                    location_type='comment',
                                    location_ids=(cmt_idx, para_idx),
                                    runs_info=[RunInfo(text=text)],
                                    full_text=text
                                ))
                                stats['comment'] += 1
        except Exception:
            pass

        # 9. 도형(Shape) 내 대체 텍스트 수집
        stats['shape'] = 0
        try:
            # DrawingML 네임스페이스
            wp_ns = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
            a_ns = 'http://schemas.openxmlformats.org/drawingml/2006/main'

            body = doc.element.body
            # 모든 drawing 요소에서 텍스트 추출
            for shape_idx, drawing in enumerate(body.iter('{%s}drawing' % (
                'http://schemas.openxmlformats.org/wordprocessingml/2006/main'))):
                # 도형 내 텍스트 바디 찾기
                for txBody in drawing.iter('{%s}txBody' % a_ns):
                    for para_idx, para_elem in enumerate(txBody.findall('{%s}p' % a_ns)):
                        texts = []
                        for t_elem in para_elem.iter('{%s}t' % a_ns):
                            if t_elem.text:
                                texts.append(t_elem.text)
                        text = ''.join(texts)
                        if text.strip():
                            units.append(TextUnit(
                                location_type='shape',
                                location_ids=(shape_idx, para_idx),
                                runs_info=[RunInfo(text=text)],
                                full_text=text
                            ))
                            stats['shape'] += 1
        except Exception:
            pass

        # 10. SmartArt 텍스트 수집
        stats['smartart'] = 0
        try:
            dgm_ns = 'http://schemas.openxmlformats.org/drawingml/2006/diagram'
            a_ns = 'http://schemas.openxmlformats.org/drawingml/2006/main'

            body = doc.element.body
            for smart_idx, dgm in enumerate(body.iter('{%s}relIds' % dgm_ns)):
                # SmartArt 데이터에서 텍스트 추출
                parent = dgm.getparent()
                if parent is not None:
                    for txBody in parent.iter('{%s}txBody' % a_ns):
                        for para_idx, para_elem in enumerate(txBody.findall('{%s}p' % a_ns)):
                            texts = []
                            for t_elem in para_elem.iter('{%s}t' % a_ns):
                                if t_elem.text:
                                    texts.append(t_elem.text)
                            text = ''.join(texts)
                            if text.strip():
                                units.append(TextUnit(
                                    location_type='smartart',
                                    location_ids=(smart_idx, para_idx),
                                    runs_info=[RunInfo(text=text)],
                                    full_text=text
                                ))
                                stats['smartart'] += 1
        except Exception:
            pass

        # 11. 하이퍼링크 텍스트 수집 (별도 run에 있는 경우)
        stats['hyperlink'] = 0
        try:
            body = doc.element.body
            for hl_idx, hyperlink in enumerate(body.iter(qn('w:hyperlink'))):
                # 하이퍼링크 내 텍스트 추출
                text = ''.join(t.text or '' for t in hyperlink.iter(qn('w:t')))
                if text.strip():
                    # 본문 단락에서 이미 수집된 텍스트인지 확인
                    already_collected = any(
                        u.full_text == text and u.location_type == 'paragraph'
                        for u in units
                    )
                    if not already_collected:
                        units.append(TextUnit(
                            location_type='hyperlink',
                            location_ids=(hl_idx, 0),
                            runs_info=[RunInfo(text=text)],
                            full_text=text
                        ))
                        stats['hyperlink'] += 1
        except Exception:
            pass

        # 12. 목차(Table of Contents) 필드 수집
        stats['toc'] = 0
        try:
            body = doc.element.body
            for fld_idx, fld_simple in enumerate(body.iter(qn('w:fldSimple'))):
                instr = fld_simple.get(qn('w:instr'))
                if instr and 'TOC' in instr.upper():
                    text = ''.join(t.text or '' for t in fld_simple.iter(qn('w:t')))
                    if text.strip():
                        units.append(TextUnit(
                            location_type='toc',
                            location_ids=(fld_idx, 0),
                            runs_info=[RunInfo(text=text)],
                            full_text=text
                        ))
                        stats['toc'] += 1
        except Exception:
            pass

        # 수집 통계 로깅
        total = sum(stats.values())
        self.log(f"텍스트 수집 완료 (총 {total}개): 본문 {stats['paragraph']}, 표 {stats['table']}, "
                 f"헤더 {stats['header']}, 푸터 {stats['footer']}, SDT {stats['sdt']}, "
                 f"텍스트박스 {stats['textbox']}, 각주 {stats['footnote']}, 미주 {stats['endnote']}, "
                 f"주석 {stats['comment']}, 도형 {stats['shape']}, SmartArt {stats['smartart']}, "
                 f"하이퍼링크 {stats['hyperlink']}, 목차 {stats['toc']}")

        return units

    def _collect_table_text(
        self,
        table,
        table_id,
        units: List[TextUnit],
        stats: dict,
        location_prefix: str = 'table'
    ):
        """표에서 텍스트 수집 (중첩 표 포함)"""
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                # 셀 내 단락
                for para_idx, para in enumerate(cell.paragraphs):
                    unit = self._extract_paragraph_unit(
                        para, location_prefix, (table_id, row_idx, cell_idx, para_idx)
                    )
                    if unit:
                        units.append(unit)
                        stats['table'] += 1

                # 셀 내 중첩 표
                for nested_idx, nested_table in enumerate(cell.tables):
                    self._collect_table_text(
                        nested_table,
                        f"{table_id}_nested_{nested_idx}",
                        units, stats, location_prefix
                    )

    def _extract_paragraph_unit(
        self,
        para,
        location_type: str,
        location_ids: tuple
    ) -> Optional[TextUnit]:
        """단락에서 TextUnit 추출"""
        if not para.text.strip():
            return None

        runs_info = []
        for run in para.runs:
            if run.text:  # 빈 run도 스타일 정보 유지를 위해 수집
                info = self._extract_run_info(run)
                runs_info.append(info)

        if not runs_info:
            return None

        full_text = para.text
        return TextUnit(
            location_type=location_type,
            location_ids=location_ids,
            runs_info=runs_info,
            full_text=full_text
        )

    def _extract_run_info(self, run) -> RunInfo:
        """Run에서 스타일 정보 추출"""
        font_color_rgb = None
        try:
            if run.font.color and run.font.color.rgb:
                rgb = run.font.color.rgb
                font_color_rgb = (rgb[0], rgb[1], rgb[2])
        except Exception:
            pass

        font_size = None
        try:
            if run.font.size:
                font_size = run.font.size.pt
        except Exception:
            pass

        return RunInfo(
            text=run.text,
            bold=run.bold,
            italic=run.italic,
            underline=run.underline is not None and run.underline != False,
            strike=run.font.strike,
            font_name=run.font.name,
            font_size=font_size,
            font_color_rgb=font_color_rgb,
            highlight_color=str(run.font.highlight_color) if run.font.highlight_color else None
        )

    def _translate_with_dynamic_batch(self, units: List[TextUnit]) -> List[str]:
        """동적 배치 크기로 번역"""
        translations = []
        total = len(units)

        # 텍스트 길이에 따라 그룹화
        short_texts = []  # < 100자
        medium_texts = []  # 100-500자
        long_texts = []  # > 500자

        for i, unit in enumerate(units):
            text_len = len(unit.full_text)
            if text_len < 100:
                short_texts.append((i, unit.full_text))
            elif text_len < 500:
                medium_texts.append((i, unit.full_text))
            else:
                long_texts.append((i, unit.full_text))

        # 결과 저장용 딕셔너리
        result_map = {}

        # 짧은 텍스트 배치 처리
        if short_texts:
            self.log(f"짧은 텍스트 번역: {len(short_texts)}개")
            self._batch_translate(short_texts, result_map, self.BATCH_SIZE_SMALL)

        # 중간 텍스트 배치 처리
        if medium_texts:
            self.log(f"중간 텍스트 번역: {len(medium_texts)}개")
            self._batch_translate(medium_texts, result_map, self.BATCH_SIZE_MEDIUM)

        # 긴 텍스트 배치 처리
        if long_texts:
            self.log(f"긴 텍스트 번역: {len(long_texts)}개")
            self._batch_translate(long_texts, result_map, self.BATCH_SIZE_LARGE)

        # 원래 순서대로 결과 정렬
        for i in range(total):
            translations.append(result_map.get(i, units[i].full_text))

        return translations

    def _batch_translate(
        self,
        indexed_texts: List[Tuple[int, str]],
        result_map: Dict[int, str],
        batch_size: int
    ):
        """배치 번역 수행"""
        for i in range(0, len(indexed_texts), batch_size):
            batch = indexed_texts[i:i + batch_size]
            texts = [t[1] for t in batch]
            indices = [t[0] for t in batch]

            try:
                translated = self.translate_texts(texts)
                for idx, trans in zip(indices, translated):
                    result_map[idx] = trans
            except Exception as e:
                self.log(f"배치 번역 오류, 개별 번역 시도: {e}")
                # 개별 번역 폴백
                for idx, text in batch:
                    try:
                        result_map[idx] = self.translate_texts([text])[0]
                    except Exception:
                        result_map[idx] = text  # 실패 시 원본 유지

            self.log(f"진행: {min(i + batch_size, len(indexed_texts))}/{len(indexed_texts)}")

    def _apply_translations(
        self,
        doc,
        units: List[TextUnit],
        translations: List[str]
    ) -> int:
        """번역 결과 적용 - 스타일 보존"""
        from docx.shared import Pt, RGBColor

        translated_count = 0

        for unit, translated in zip(units, translations):
            try:
                para = self._get_paragraph(doc, unit)
                if para is None:
                    continue

                # XML 기반 요소 처리
                if isinstance(para, tuple) and para[0] == 'xml_element':
                    success = self._apply_xml_translation(doc, para[1], para[2], translated)
                    if success:
                        translated_count += 1
                    continue

                # 번역 텍스트를 Run들에 분배
                self._distribute_text_to_runs(para, unit.runs_info, translated)
                translated_count += 1

            except Exception as e:
                self.log(f"적용 실패 ({unit.location_type}): {e}")

        return translated_count

    def _apply_xml_translation(
        self,
        doc,
        element_type: str,
        location_ids: tuple,
        translated: str
    ) -> bool:
        """XML 기반 요소에 번역 적용 (각주, 미주, 텍스트박스)"""
        try:
            from docx.oxml.ns import qn

            if element_type == 'footnote':
                fn_idx, para_idx = location_ids
                for rel in doc.part.rels.values():
                    if 'footnotes' in rel.reltype:
                        footnotes_xml = rel.target_part.element
                        footnotes = list(footnotes_xml.findall(qn('w:footnote')))
                        if fn_idx < len(footnotes):
                            paras = list(footnotes[fn_idx].findall(qn('w:p')))
                            if para_idx < len(paras):
                                self._replace_paragraph_text_xml(paras[para_idx], translated)
                                return True

            elif element_type == 'endnote':
                en_idx, para_idx = location_ids
                for rel in doc.part.rels.values():
                    if 'endnotes' in rel.reltype:
                        endnotes_xml = rel.target_part.element
                        endnotes = list(endnotes_xml.findall(qn('w:endnote')))
                        if en_idx < len(endnotes):
                            paras = list(endnotes[en_idx].findall(qn('w:p')))
                            if para_idx < len(paras):
                                self._replace_paragraph_text_xml(paras[para_idx], translated)
                                return True

            elif element_type == 'textbox':
                txbx_idx, para_idx = location_ids
                body = doc.element.body
                textboxes = list(body.iter(qn('w:txbxContent')))
                if txbx_idx < len(textboxes):
                    paras = list(textboxes[txbx_idx].findall(qn('w:p')))
                    if para_idx < len(paras):
                        self._replace_paragraph_text_xml(paras[para_idx], translated)
                        return True

            elif element_type == 'sdt':
                sdt_idx, para_idx = location_ids
                body = doc.element.body
                sdts = list(body.iter(qn('w:sdt')))
                if sdt_idx < len(sdts):
                    sdt_content = sdts[sdt_idx].find(qn('w:sdtContent'))
                    if sdt_content is not None:
                        paras = list(sdt_content.findall(qn('w:p')))
                        if para_idx < len(paras):
                            self._replace_paragraph_text_xml(paras[para_idx], translated)
                            return True

            elif element_type == 'comment':
                cmt_idx, para_idx = location_ids
                for rel in doc.part.rels.values():
                    if 'comments' in rel.reltype:
                        comments_xml = rel.target_part.element
                        comments = list(comments_xml.findall(qn('w:comment')))
                        if cmt_idx < len(comments):
                            paras = list(comments[cmt_idx].findall(qn('w:p')))
                            if para_idx < len(paras):
                                self._replace_paragraph_text_xml(paras[para_idx], translated)
                                return True

            elif element_type == 'shape':
                shape_idx, para_idx = location_ids
                a_ns = 'http://schemas.openxmlformats.org/drawingml/2006/main'
                body = doc.element.body
                drawings = list(body.iter('{%s}drawing' % (
                    'http://schemas.openxmlformats.org/wordprocessingml/2006/main')))
                if shape_idx < len(drawings):
                    txBodies = list(drawings[shape_idx].iter('{%s}txBody' % a_ns))
                    if txBodies:
                        paras = list(txBodies[0].findall('{%s}p' % a_ns))
                        if para_idx < len(paras):
                            self._replace_drawingml_text(paras[para_idx], translated, a_ns)
                            return True

            elif element_type == 'smartart':
                smart_idx, para_idx = location_ids
                dgm_ns = 'http://schemas.openxmlformats.org/drawingml/2006/diagram'
                a_ns = 'http://schemas.openxmlformats.org/drawingml/2006/main'
                body = doc.element.body
                dgm_elements = list(body.iter('{%s}relIds' % dgm_ns))
                if smart_idx < len(dgm_elements):
                    parent = dgm_elements[smart_idx].getparent()
                    if parent is not None:
                        txBodies = list(parent.iter('{%s}txBody' % a_ns))
                        if txBodies:
                            paras = list(txBodies[0].findall('{%s}p' % a_ns))
                            if para_idx < len(paras):
                                self._replace_drawingml_text(paras[para_idx], translated, a_ns)
                                return True

            elif element_type == 'hyperlink':
                hl_idx, para_idx = location_ids
                body = doc.element.body
                hyperlinks = list(body.iter(qn('w:hyperlink')))
                if hl_idx < len(hyperlinks):
                    paras = list(hyperlinks[hl_idx].iter(qn('w:r')))
                    for r_elem in paras:
                        text_elements = list(r_elem.iter(qn('w:t')))
                        if text_elements:
                            text_elements[0].text = translated
                            for t_elem in text_elements[1:]:
                                t_elem.text = ""
                            return True

            elif element_type == 'toc':
                fld_idx, _ = location_ids
                body = doc.element.body
                fld_simples = list(body.iter(qn('w:fldSimple')))
                if fld_idx < len(fld_simples):
                    text_elements = list(fld_simples[fld_idx].iter(qn('w:t')))
                    if text_elements:
                        text_elements[0].text = translated
                        for t_elem in text_elements[1:]:
                            t_elem.text = ""
                        return True

        except Exception as e:
            self.log(f"XML 번역 적용 실패 ({element_type}): {e}")

        return False

    def _replace_drawingml_text(self, para_elem, new_text: str, a_ns: str):
        """DrawingML 단락의 텍스트 교체 (도형, SmartArt용)"""
        # 기존 텍스트 요소 찾기
        text_elements = list(para_elem.iter('{%s}t' % a_ns))
        if not text_elements:
            return

        # 첫 번째 텍스트 요소에 전체 번역 텍스트 설정
        text_elements[0].text = new_text

        # 나머지 텍스트 요소는 비우기
        for t_elem in text_elements[1:]:
            t_elem.text = ""

    def _replace_paragraph_text_xml(self, para_elem, new_text: str):
        """XML 단락의 텍스트 교체"""
        from docx.oxml.ns import qn

        # 기존 텍스트 요소 찾기
        text_elements = list(para_elem.iter(qn('w:t')))
        if not text_elements:
            return

        # 첫 번째 텍스트 요소에 전체 번역 텍스트 설정
        text_elements[0].text = new_text

        # 나머지 텍스트 요소는 비우기
        for t_elem in text_elements[1:]:
            t_elem.text = ""

    def _get_paragraph(self, doc, unit: TextUnit):
        """위치 정보로 단락 가져오기"""
        try:
            if unit.location_type == 'paragraph':
                para_idx = unit.location_ids[0]
                return doc.paragraphs[para_idx]

            elif unit.location_type == 'table':
                table_id, row_idx, cell_idx, para_idx = unit.location_ids
                # 중첩 표 처리 (table_id가 문자열일 수 있음)
                if isinstance(table_id, int):
                    return doc.tables[table_id].rows[row_idx].cells[cell_idx].paragraphs[para_idx]
                else:
                    # 중첩 표는 XML로 처리
                    return ('xml_element', unit.location_type, unit.location_ids)

            elif unit.location_type == 'header':
                # 새 형식: (section_idx, header_type, para_idx) 또는 기존 형식
                if len(unit.location_ids) == 3:
                    section_idx, header_type, para_idx = unit.location_ids
                    header = getattr(doc.sections[section_idx], header_type, None)
                    if header:
                        return header.paragraphs[para_idx]
                else:
                    section_idx, para_idx = unit.location_ids
                    return doc.sections[section_idx].header.paragraphs[para_idx]

            elif unit.location_type == 'footer':
                # 새 형식: (section_idx, footer_type, para_idx) 또는 기존 형식
                if len(unit.location_ids) == 3:
                    section_idx, footer_type, para_idx = unit.location_ids
                    footer = getattr(doc.sections[section_idx], footer_type, None)
                    if footer:
                        return footer.paragraphs[para_idx]
                else:
                    section_idx, para_idx = unit.location_ids
                    return doc.sections[section_idx].footer.paragraphs[para_idx]

            elif unit.location_type in ('footnote', 'endnote', 'textbox', 'sdt',
                                        'comment', 'shape', 'smartart', 'hyperlink', 'toc'):
                # XML 기반 요소는 특별 처리 필요
                return ('xml_element', unit.location_type, unit.location_ids)

            elif unit.location_type in ('header_table', 'footer_table'):
                # 헤더/푸터 내 표는 XML로 처리
                return ('xml_element', unit.location_type, unit.location_ids)

        except (IndexError, KeyError, AttributeError):
            return None

        return None

    def _distribute_text_to_runs(
        self,
        para,
        original_runs_info: List[RunInfo],
        translated_text: str
    ):
        """번역된 텍스트를 원본 Run 구조에 맞게 분배"""
        from docx.shared import Pt, RGBColor

        if not para.runs:
            para.text = translated_text
            return

        # 원본 텍스트 총 길이
        original_total_len = sum(len(info.text) for info in original_runs_info)
        if original_total_len == 0:
            para.runs[0].text = translated_text
            return

        # Run이 하나뿐이면 단순 교체
        if len(para.runs) == 1:
            self._apply_style_to_run(para.runs[0], original_runs_info[0])
            para.runs[0].text = translated_text
            return

        # 여러 Run의 경우: 비율에 따라 텍스트 분배
        translated_len = len(translated_text)
        current_pos = 0

        for i, (run, info) in enumerate(zip(para.runs, original_runs_info)):
            # 스타일 적용
            self._apply_style_to_run(run, info)

            # 마지막 run이면 나머지 텍스트 전부 할당
            if i == len(para.runs) - 1:
                run.text = translated_text[current_pos:]
            else:
                # 비율에 따라 텍스트 분배
                ratio = len(info.text) / original_total_len
                chars_for_run = max(1, int(translated_len * ratio))

                # 단어 경계에서 자르기 시도
                end_pos = min(current_pos + chars_for_run, translated_len)

                # 단어 중간이면 공백까지 확장
                if end_pos < translated_len and translated_text[end_pos] not in ' \t\n':
                    # 앞으로 공백 찾기
                    space_pos = translated_text.find(' ', end_pos)
                    if space_pos != -1 and space_pos - end_pos < 10:
                        end_pos = space_pos + 1

                run.text = translated_text[current_pos:end_pos]
                current_pos = end_pos

        # 남은 run들 비우기
        for run in para.runs[len(original_runs_info):]:
            run.text = ""

    def _apply_style_to_run(self, run, info: RunInfo):
        """Run에 스타일 적용"""
        from docx.shared import Pt, RGBColor

        try:
            if info.bold is not None:
                run.bold = info.bold
            if info.italic is not None:
                run.italic = info.italic
            if info.underline:
                run.underline = True
            if info.strike is not None:
                run.font.strike = info.strike
            if info.font_name:
                run.font.name = info.font_name
            if info.font_size:
                run.font.size = Pt(info.font_size)
            if info.font_color_rgb:
                run.font.color.rgb = RGBColor(*info.font_color_rgb)
        except Exception:
            pass  # 스타일 적용 실패해도 계속 진행
